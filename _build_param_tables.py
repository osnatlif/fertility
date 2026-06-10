"""Build a Word document with two tables:
  Table 1: Model parameters by SOURCE FILE, with an Estimation group column
           A. Main parameters file (parameters.py)                       38
           B. Cohort file -- employment hazards                          24  (Osnat: 18 + 6 exp^2)
           C. Cohort file -- marriage offer                              5
           D. Cohort file -- wage process                                26
  Table 2: Setup constants (institutional/lifecycle: UB, childcare, ...)

Output: parameters_table.docx in this directory.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ----------------------------------------------------------------------
# Data
# ----------------------------------------------------------------------
# Each row: (estimation_group, param_name, definition)
#   estimation_group : phase(s) that move this param, or "FIXED (...)"

# ===== Section A: Main parameters file (parameters.py) -- 38 =====
SEC_A_MAIN = [
    # marriage utility (4)
    ("marriage, all",                  "mconst1_h",        "Husband's utility from a 1-step educational gap (hypergamy)"),
    ("marriage, all",                  "mconst1_w",        "Wife's utility from a 1-step educational gap"),
    ("marriage, all",                  "mconst2_h",        "Husband's utility from a 2-step educational gap"),
    ("marriage, all",                  "mconst2_w",        "Wife's utility from a 2-step educational gap"),
    # marriage homogamy taste (3)
    ("marriage, all",                  "taste_hs",         "Taste for homogamous marriage at HS education level"),
    ("marriage, all",                  "taste_sc",         "Taste for homogamous marriage at Some College level"),
    ("marriage, all",                  "taste_cg",         "Taste for homogamous marriage at College Graduate+ level"),
    # utility (10)
    ("FIXED (EKL scale)",              "alpha0",           "CRRA consumption curvature (scale parameter)"),
    ("marriage, employment, all",      "alpha2w0",         "Wife leisure utility constant"),
    ("marriage, employment, all",      "alpha2w1",         "Wife leisure utility × number of kids"),
    ("employment, all",                "alpha2h0",         "Husband leisure utility constant"),
    ("employment, all",                "alpha2h1",         "Husband leisure utility × number of kids"),
    ("marriage, fertility, all",       "alpha3_w_m",       "Wife's utility from children when married"),
    ("marriage, fertility, all",       "alpha3_w_s",       "Wife's utility from children when single"),
    ("marriage, fertility, all",       "alpha3_h_m",       "Husband's utility from children when married"),
    ("fertility, all",                 "alpha3_h_s",       "Husband's utility from children when single"),
    ("fertility, all",                 "alpha4",           "Curvature of kids quality utility"),
    # marriage / divorce cost (3)
    ("marriage, all",                  "mc",               "Fixed cost of getting married"),
    ("marriage, all",                  "dc_w",             "Divorce cost — wife (utility penalty of divorce)"),
    ("marriage, all",                  "dc_h",             "Divorce cost — husband"),
    # variances (6)
    ("FIXED (external)",               "sigma_ability_w",  "Wife ability shock standard deviation"),
    ("FIXED (external)",               "sigma_ability_h",  "Husband ability shock standard deviation"),
    ("wage, all",                      "sigma_w_wage",     "Wife log-wage residual standard deviation"),
    ("wage, all",                      "sigma_h_wage",     "Husband log-wage residual standard deviation"),
    ("marriage, all",                  "sigma_q",          "Match-quality shock standard deviation"),
    ("fertility, all",                 "sigma_p",          "Pregnancy-preference shock standard deviation"),
    # terminal values (12 = 4 fixed + 2 estimated, each × 2 genders)
    ("FIXED (weak ID)",                "t1_w",             "Terminal value — wife: own SC education"),
    ("FIXED (weak ID)",                "t2_w",             "Terminal value — wife: own CG+ education"),
    ("FIXED (weak ID)",                "t3_w",             "Terminal value — wife: husband's SC education"),
    ("FIXED (weak ID)",                "t4_w",             "Terminal value — wife: husband's CG+ education"),
    ("marriage, all",                  "t5_w",             "Terminal value — wife: marriage utility at retirement"),
    ("marriage, all",                  "t6_w",             "Terminal value — wife: utility from children at retirement"),
    ("FIXED (weak ID)",                "t1_h",             "Terminal value — husband: wife's SC education"),
    ("FIXED (weak ID)",                "t2_h",             "Terminal value — husband: wife's CG+ education"),
    ("FIXED (weak ID)",                "t3_h",             "Terminal value — husband: own SC education"),
    ("FIXED (weak ID)",                "t4_h",             "Terminal value — husband: own CG+ education"),
    ("marriage, all",                  "t5_h",             "Terminal value — husband: marriage utility at retirement"),
    ("marriage, all",                  "t6_h",             "Terminal value — husband: utility from children at retirement"),
]   # 38 rows


# ===== Section B: Cohort file -- employment hazards -- 24 (Osnat's count: 18) =====
SEC_B_EMP = [
    # Wife full-time offer hazard (4)
    ("employment, all",  "lambda0_w_ft",    "Wife full-time job-offer probability — constant"),
    ("employment, all",  "lambda1_w_ft",    "Wife full-time job-offer probability — experience slope"),
    ("employment, all",  "lambda2_w_ft",    "Wife full-time job-offer probability — education slope"),
    ("employment, all",  "lambda15_w_ft",   "Wife full-time job-offer probability — experience² (capped at 0; concave-only)"),
    # Wife part-time offer hazard (4)
    ("employment, all",  "lambda0_w_pt",    "Wife part-time job-offer probability — constant"),
    ("employment, all",  "lambda1_w_pt",    "Wife part-time job-offer probability — experience slope"),
    ("employment, all",  "lambda2_w_pt",    "Wife part-time job-offer probability — education slope"),
    ("employment, all",  "lambda15_w_pt",   "Wife part-time job-offer probability — experience² (capped at 0)"),
    # Wife not-laid-off hazard (4)
    ("employment, all",  "lambda0_w_f",     "Wife not-laid-off probability — constant"),
    ("employment, all",  "lambda1_w_f",     "Wife not-laid-off probability — experience slope"),
    ("employment, all",  "lambda2_w_f",     "Wife not-laid-off probability — education slope"),
    ("employment, all",  "lambda15_w_f",    "Wife not-laid-off probability — experience² (capped at 0)"),
    # Husband full-time offer hazard (4) -- all FIXED
    ("FIXED (external)", "lambda0_h_ft",    "Husband full-time job-offer probability — constant"),
    ("FIXED (external)", "lambda1_h_ft",    "Husband full-time job-offer probability — experience slope"),
    ("FIXED (external)", "lambda2_h_ft",    "Husband full-time job-offer probability — education slope"),
    ("FIXED (external)", "lambda15_h_ft",   "Husband full-time job-offer probability — experience²"),
    # Husband part-time offer hazard (4) -- all FIXED
    ("FIXED (external)", "lambda0_h_pt",    "Husband part-time job-offer probability — constant"),
    ("FIXED (external)", "lambda1_h_pt",    "Husband part-time job-offer probability — experience slope"),
    ("FIXED (external)", "lambda2_h_pt",    "Husband part-time job-offer probability — education slope"),
    ("FIXED (external)", "lambda15_h_pt",   "Husband part-time job-offer probability — experience²"),
    # Husband not-laid-off hazard (4) -- all FIXED
    ("FIXED (external)", "lambda0_h_f",     "Husband not-laid-off probability — constant"),
    ("FIXED (external)", "lambda1_h_f",     "Husband not-laid-off probability — experience slope"),
    ("FIXED (external)", "lambda2_h_f",     "Husband not-laid-off probability — education slope"),
    ("FIXED (external)", "lambda15_h_f",    "Husband not-laid-off probability — experience²"),
]   # 24 rows total; Osnat's count of 18 = excludes the 6 lambda15 exp² terms


# ===== Section C: Cohort file -- marriage offer -- 5 =====
SEC_C_MAR = [
    ("marriage, all",     "omega3",   "Meeting-probability intercept (shared by both genders)"),
    ("marriage, all",     "omega4_w", "Wife meeting probability — age slope (≥0; peak ≈30, ~0 at 48)"),
    ("marriage, all",     "omega5_w", "Wife meeting probability — age² slope (≤0, concave)"),
    ("marriage, all",     "omega4_h", "Husband meeting probability — age slope (≥0)"),
    ("marriage, all",     "omega5_h", "Husband meeting probability — age² slope (≤0)"),
]   # 5 rows


# ===== Section D: Cohort file -- wage process -- 26 =====
SEC_D_WAGE = [
    # Wife (13)
    ("wage",  "beta0_w",   "Wife log-wage: ability slope"),
    ("wage",  "beta11_w",  "Wife log-wage: experience × HS"),
    ("wage",  "beta12_w",  "Wife log-wage: experience × SC"),
    ("wage",  "beta13_w",  "Wife log-wage: experience × CG+"),
    ("wage",  "beta21_w",  "Wife log-wage: experience² × HS"),
    ("wage",  "beta22_w",  "Wife log-wage: experience² × SC"),
    ("wage",  "beta23_w",  "Wife log-wage: experience² × CG+"),
    ("wage",  "beta31_w",  "Wife log-wage: intercept × HS"),
    ("wage",  "beta32_w",  "Wife log-wage: intercept × SC"),
    ("wage",  "beta33_w",  "Wife log-wage: intercept × CG+"),
    ("wage",  "beta41_w",  "Wife log-wage: not-employed-previous-period × HS"),
    ("wage",  "beta42_w",  "Wife log-wage: not-employed-previous-period × SC"),
    ("wage",  "beta43_w",  "Wife log-wage: not-employed-previous-period × CG+"),
    # Husband (13)
    ("wage",  "beta0_h",   "Husband log-wage: ability slope"),
    ("wage",  "beta11_h",  "Husband log-wage: experience × HS"),
    ("wage",  "beta12_h",  "Husband log-wage: experience × SC"),
    ("wage",  "beta13_h",  "Husband log-wage: experience × CG+"),
    ("wage",  "beta21_h",  "Husband log-wage: experience² × HS"),
    ("wage",  "beta22_h",  "Husband log-wage: experience² × SC"),
    ("wage",  "beta23_h",  "Husband log-wage: experience² × CG+"),
    ("wage",  "beta31_h",  "Husband log-wage: intercept × HS"),
    ("wage",  "beta32_h",  "Husband log-wage: intercept × SC"),
    ("wage",  "beta33_h",  "Husband log-wage: intercept × CG+"),
    ("wage",  "beta41_h",  "Husband log-wage: not-employed-previous-period × HS"),
    ("wage",  "beta42_h",  "Husband log-wage: not-employed-previous-period × SC"),
    ("wage",  "beta43_h",  "Husband log-wage: not-employed-previous-period × CG+"),
]   # 26 rows


# ===== Table 2: Setup constants (NOT model parameters) =====
# These live in constant_parameters.pyx (institutional / lifecycle structure).
CONSTANTS = [
    ("beta0 (discount)", "Annual discount factor",                                 "0.983"),
    ("scale",            "Equivalence scale (fraction of public consumption)",     "0.707"),
    ("bp",               "Intra-household bargaining power",                       "0.5"),
    ("eta1",             "Kids' share of household income — 1 kid",           "0.194"),
    ("eta2",             "Kids' share of household income — 2 kids",          "0.293"),
    ("eta3",             "Kids' share of household income — 3 kids",          "0.367"),
    ("eta4",             "Kids' share of household income — 4 kids",          "0.423"),
    ("ub_w",             "Unemployment benefit — wife (annual, USD)",         "6,000"),
    ("ub_h",             "Unemployment benefit — husband (annual, USD)",      "8,000"),
    ("childcare_cost",   "Childcare cost per child below age 5 (annual, USD)",     "5,000"),
    ("cb_const",         "Child benefit base — single mom + 1 kid (1960, annual USD)", "3,000"),
    ("cb_per_child",     "Child benefit per additional kid (1960, annual USD)",    "1,517.235"),
    ("max_period",       "Periods in the lifecycle (t=1..33 covers ages 18–50)", "34"),
    ("AGE_VALUES",       "Education entry ages (HS / SC / CG+)",                   "{18, 20, 22}"),
    ("MAX_FERTILITY_AGE","Maximum fertility age",                                  "45"),
    ("N_GH",             "Gauss-Hermite quadrature nodes (match quality)",         "3"),
    ("N_GH_PREG",        "Gauss-Hermite quadrature nodes (pregnancy)",             "3"),
    ("N_EXP",            "Wife experience grid size",                              "4"),
    ("exp_grid",         "Wife experience grid points (years)",                    "{1, 4, 8, 12}"),
    ("DRAW_F",           "Forward simulation draws (sample size)",                 "5,000"),
    ("preg_married",     "Utility shock — marital × pregnancy interaction (pinned at 0)",  "0"),
    ("preg_unmarried",   "Utility shock — single × pregnancy interaction (pinned at 0)",   "0"),
    ("preg_kids",        "Utility shock — kids × pregnancy interaction (pinned at 0)",     "0"),
]


# ----------------------------------------------------------------------
# Build helpers
# ----------------------------------------------------------------------

def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text, bold=False, size=9, color=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color is not None:
        run.font.color.rgb = color


def _set_table_borders(tbl):
    tblPr = tbl._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._element.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '999999')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _row_shade_for(group):
    """Color the row by the first phase in the estimation group (or gray if FIXED)."""
    if group.startswith("FIXED"):
        return "F2F2F2"
    primary = group.split(",")[0].strip()
    return {
        "marriage":   "DDEBF7",
        "employment": "FFF2CC",
        "fertility":  "E2EFDA",
        "wage":       "FCE4D6",
        "all":        "FFFFFF",
    }.get(primary, "FFFFFF")


def add_param_subtable(doc, title, rows, note=None):
    """Add a section heading + a 4-column table for one source-file block."""
    h = doc.add_heading(title, level=2)
    for run in h.runs:
        run.font.name = "Calibri"

    tbl = doc.add_table(rows=1 + len(rows), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(tbl)

    col_widths = [Inches(1.6), Inches(1.3), Inches(2.6), Inches(1.0)]
    for row in tbl.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]

    headers = ["Estimation group(s)", "Parameter", "Definition", "Estimated value"]
    for idx, h in enumerate(headers):
        cell = tbl.rows[0].cells[idx]
        _set_cell_text(cell, h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade_cell(cell, "305496")

    for r_idx, (group, name, defn) in enumerate(rows, start=1):
        cells = [group, name, defn, ""]
        for c_idx, value in enumerate(cells):
            cell = tbl.rows[r_idx].cells[c_idx]
            _set_cell_text(cell, value, bold=(c_idx == 1), size=9)
            _shade_cell(cell, _row_shade_for(group))

    if note:
        p = doc.add_paragraph()
        r = p.add_run(note)
        r.italic = True
        r.font.size = Pt(9)
        r.font.name = "Calibri"


def build_doc(output_path):
    doc = Document()

    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin  = section.bottom_margin = Inches(1)

    # Title
    title = doc.add_heading("Structural Model Parameters", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("1960 white cohort — household model of marriage, fertility, female labor supply")
    r.italic = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    # Overview line
    ov = doc.add_paragraph()
    ov.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ov.add_run(
        "Sections A + B + C = 61 structural parameters "
        "(38 main file + 18 employment lambdas + 5 marriage offer; "
        "Section B includes 6 additional exp² terms). Section D = 26 wage parameters "
        "(external pre-estimates; sigmas in A). The Estimation group column lists which "
        "phase(s) move each parameter; FIXED rows are not estimated."
    )
    r.font.size = Pt(9)
    r.font.name = "Calibri"

    doc.add_paragraph()

    # Table 1 banner
    h1 = doc.add_heading("Table 1. Model Parameters by Source File", level=1)
    for run in h1.runs:
        run.font.name = "Calibri"

    # Section A
    add_param_subtable(
        doc,
        "A. Main parameters file (parameters.py) — 38 params",
        SEC_A_MAIN,
        note="Counts: 38 attributes set on `p` in parameters.py "
             "(41 total minus the 3 preg_* shocks pinned at 0, which appear in Table 2)."
    )

    doc.add_paragraph()

    # Section B
    add_param_subtable(
        doc,
        "B. Cohort file (input/parameters1960white.py) — employment hazards — 24 params",
        SEC_B_EMP,
        note="Osnat's count of 18 = lambda{0,1,2} × {FT,PT,fired} × {wife,husband} "
             "(3×3×2). The 6 lambda15 (experience²) terms above are additional; "
             "the 3 wife lambda15 are activated by the employment phase (capped at 0 = "
             "concave-only); the 3 husband lambda15 stay at external values "
             "(lambda15_h_ft = −0.0015; others = 0)."
    )

    doc.add_page_break()

    # Section C
    add_param_subtable(
        doc,
        "C. Cohort file (input/parameters1960white.py) — marriage offer — 5 params",
        SEC_C_MAR,
        note="Until 2026-06-10 the model used a single meeting-probability curve "
             "(wife's omegas) for both sexes — a bug. Fixed: meeting_partner.pyx now "
             "exposes prob_w(age) and prob_h(age); the husband age slopes are "
             "initialised equal to the wife's in the cohort file but estimated "
             "independently by the marriage / all phases. Sign restrictions: "
             "omega4_* ≥ 0 (more age → more meetings up to the peak); "
             "omega5_* ≤ 0 (concave so meetings decline at later ages)."
    )

    doc.add_paragraph()

    # Section D
    add_param_subtable(
        doc,
        "D. Cohort file (input/parameters1960white.py) — wage process — 26 params",
        SEC_D_WAGE,
        note="Wage betas are externally pre-estimated and excluded from the `all` phase. "
             "The wage phase moves them only if you explicitly select it; identification "
             "is weak because wage moments are NOT in the structural objective "
             "(diagnostics only). The 2 wage shock SDs (sigma_w_wage, sigma_h_wage) "
             "live in parameters.py and appear in Section A."
    )

    doc.add_page_break()

    # ---------- Table 2: Setup constants ----------
    h2 = doc.add_heading("Table 2. Setup Constants (institutional & lifecycle structure)", level=1)
    for run in h2.runs:
        run.font.name = "Calibri"

    tbl2 = doc.add_table(rows=1 + len(CONSTANTS), cols=3)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(tbl2)

    col_widths2 = [Inches(1.6), Inches(3.4), Inches(1.5)]
    for row in tbl2.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths2[idx]

    for idx, h in enumerate(["Constant", "Definition", "Value"]):
        cell = tbl2.rows[0].cells[idx]
        _set_cell_text(cell, h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade_cell(cell, "305496")

    for r_idx, (name, defn, value) in enumerate(CONSTANTS, start=1):
        cells = [name, defn, value]
        for c_idx, val in enumerate(cells):
            cell = tbl2.rows[r_idx].cells[c_idx]
            _set_cell_text(cell, val, bold=(c_idx == 0), size=9)
            if r_idx % 2 == 0:
                _shade_cell(cell, "F2F2F2")

    doc.add_paragraph()
    n2 = doc.add_paragraph()
    nr2 = n2.add_run(
        "These are NOT model parameters. They are institutional inputs (UB, childcare, "
        "child benefit), preference scales fixed from literature (discount, equivalence, "
        "bargaining, kids' income share), or solution-method/lifecycle structure "
        "(quadrature nodes, experience grid, ages, simulation size). The 3 preg_* shocks "
        "appear here because they are pinned at 0 by modeling choice rather than being "
        "actively estimated."
    )
    nr2.font.size = Pt(9)
    nr2.italic = True
    nr2.font.name = "Calibri"

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    here = os.path.dirname(os.path.abspath(__file__))
    out = os.path.join(here, "parameters_table.docx")
    path = build_doc(out)
    a, b, c, d = len(SEC_A_MAIN), len(SEC_B_EMP), len(SEC_C_MAR), len(SEC_D_WAGE)
    print(f"Wrote: {path}")
    print(f"Section A (main file):                 {a:3d}")
    print(f"Section B (cohort employment):         {b:3d}  (Osnat: 18 + 6 exp^2)")
    print(f"Section C (cohort marriage offer):     {c:3d}")
    print(f"Section D (cohort wage process):       {d:3d}")
    print(f"Structural (A+B[18 only]+C):           {a + 18 + c:3d}   <- matches Osnat's 61")
    print(f"All param rows in Table 1:             {a + b + c + d:3d}")
    print(f"Setup constants rows in Table 2:       {len(CONSTANTS):3d}")
